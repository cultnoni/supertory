"""Unified proof manuscript extractors: HWP (pyhwp/hwp5) + DOCX (python-docx).

Both paths produce the same structure:
  {
    text: str,                 # full extracted body (memos may still be inline)
    memos: [{location_context, memo_content}, ...],
    format: "hwp"|"docx"|"hwpx"|...,
    parser: "pyhwp"|"python-docx"|"stdlib-zip"|...,
    warnings: [str, ...],
    title: str,
  }

Falls back to document_import (stdlib zip/XML) when optional packages are missing
or when parsing fails, so the app keeps working without hard dependencies.
"""

from __future__ import annotations

import io
import re
import tempfile
import zipfile
import xml.etree.ElementTree as ET
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any

import document_import
import proof_diff

MAX_UPLOAD_BYTES = document_import.MAX_UPLOAD_BYTES


@dataclass
class ExtractedMemo:
    location_context: str
    memo_content: str

    def to_dict(self) -> dict[str, str]:
        return {
            "location_context": self.location_context,
            "memo_content": self.memo_content,
        }


@dataclass
class UnifiedProofExtract:
    """Integrated text + memo payload for AI pipeline steps 1–3."""

    text: str
    memos: list[ExtractedMemo] = field(default_factory=list)
    format: str = "unknown"
    parser: str = "unknown"
    title: str = ""
    warnings: list[str] = field(default_factory=list)

    def to_dict(self) -> dict[str, Any]:
        return {
            "text": self.text,
            "memos": [m.to_dict() for m in self.memos],
            "format": self.format,
            "parser": self.parser,
            "title": self.title,
            "warnings": list(self.warnings),
        }


def _has_module(name: str) -> bool:
    try:
        __import__(name)
        return True
    except Exception:
        return False


def parser_status() -> dict[str, Any]:
    """Report which native parsers are available."""
    return {
        "python_docx": _has_module("docx"),
        "pyhwp": _has_module("hwp5"),
        "olefile": _has_module("olefile"),
        "six": _has_module("six"),
    }


def extract_proof_document(filename: str, data: bytes) -> UnifiedProofExtract:
    """Entry: bytes + filename → unified text/memo extract."""
    if not data:
        raise ValueError("파일이 비어 있습니다.")
    if len(data) > MAX_UPLOAD_BYTES:
        raise ValueError(
            f"파일이 너무 큽니다. {MAX_UPLOAD_BYTES // (1024 * 1024)}MB 이하만 가져올 수 있어요."
        )

    name = filename or "upload.bin"
    ext = Path(name).suffix.lower()
    title = document_import.title_from_filename(name)
    warnings: list[str] = []

    if ext == ".hwp":
        return _extract_hwp(data, title=title, filename=name)
    if ext == ".docx":
        return _extract_docx_unified(data, title=title, filename=name)
    if ext == ".hwpx":
        # Prefer zip XML (stdlib); still surface memos from plain patterns
        return _extract_via_document_import(name, data, preferred_format="hwpx")

    # Other types: document_import, then peel memos from text
    return _extract_via_document_import(name, data, preferred_format=ext.lstrip(".") or "text")


def _merge_text_memos(text: str, structured: list[ExtractedMemo]) -> tuple[str, list[ExtractedMemo]]:
    cleaned, text_memos = proof_diff.extract_memos(text)
    seen = {(m.location_context, m.memo_content) for m in structured}
    merged = list(structured)
    for tm in text_memos:
        key = (tm.location_context, tm.memo_content)
        if key in seen:
            continue
        merged.append(ExtractedMemo(tm.location_context, tm.memo_content))
        seen.add(key)
    return cleaned, merged


def _extract_via_document_import(
    filename: str,
    data: bytes,
    *,
    preferred_format: str,
) -> UnifiedProofExtract:
    extracted = document_import.extract_document(filename, data)
    text, memos = _merge_text_memos(extracted.text, [])
    return UnifiedProofExtract(
        text=text,
        memos=memos,
        format=extracted.format_name or preferred_format,
        parser="document_import",
        title=extracted.title or document_import.title_from_filename(filename),
        warnings=list(extracted.warnings),
    )


# ── DOCX (python-docx preferred) ─────────────────────────────────────

def _extract_docx_unified(data: bytes, *, title: str, filename: str) -> UnifiedProofExtract:
    if _has_module("docx"):
        try:
            return _extract_docx_python_docx(data, title=title)
        except Exception as error:  # noqa: BLE001
            # Fall through to stdlib
            fallback = _extract_via_document_import(filename, data, preferred_format="docx")
            fallback.warnings.append(f"python-docx 실패 → 기본 파서 사용: {error}")
            return fallback
    result = _extract_via_document_import(filename, data, preferred_format="docx")
    result.warnings.append(
        "python-docx가 설치되지 않아 기본 DOCX 파서를 씁니다. "
        "메모(댓글) 추출이 약할 수 있어요. pip install python-docx"
    )
    return result


def _extract_docx_python_docx(data: bytes, *, title: str) -> UnifiedProofExtract:
    from docx import Document  # type: ignore
    from docx.opc.constants import RELATIONSHIP_TYPE as RT  # type: ignore

    doc = Document(io.BytesIO(data))
    paragraphs: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            paragraphs.append(t)
    # Tables as sequential paragraphs
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                paragraphs.append(line)

    body = "\n\n".join(paragraphs)
    memos = _extract_docx_comments(doc, RT)
    body2, memos2 = _merge_text_memos(body, memos)
    return UnifiedProofExtract(
        text=body2,
        memos=memos2,
        format="docx",
        parser="python-docx",
        title=title,
        warnings=[],
    )


def _extract_docx_comments(doc: Any, RT: Any) -> list[ExtractedMemo]:
    """Read Word comments part if present."""
    memos: list[ExtractedMemo] = []
    try:
        part = doc.part.part_related_by(RT.COMMENTS)
    except Exception:
        # Also try raw zip comments.xml
        return _extract_docx_comments_from_zip(doc)

    try:
        root = ET.fromstring(part.blob)
    except Exception:
        return memos

    ns = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    }
    for c in root.findall("w:comment", ns):
        author = c.get(f"{{{ns['w']}}}author") or c.get("author") or ""
        texts = []
        for t in c.iter(f"{{{ns['w']}}}t"):
            if t.text:
                texts.append(t.text)
        content = "".join(texts).strip()
        if not content:
            continue
        ctx = f"댓글 ({author})" if author else "댓글"
        memos.append(ExtractedMemo(location_context=ctx, memo_content=content))
    return memos


def _extract_docx_comments_from_zip(doc: Any) -> list[ExtractedMemo]:
    """Fallback: open package as zip for word/comments.xml."""
    memos: list[ExtractedMemo] = []
    try:
        blob = doc.part.package.blob  # may not exist
    except Exception:
        blob = None
    if not blob:
        return memos
    try:
        zf = zipfile.ZipFile(io.BytesIO(blob))
        raw = zf.read("word/comments.xml")
    except Exception:
        return memos
    try:
        root = ET.fromstring(raw)
    except ET.ParseError:
        return memos
    ns_uri = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    for c in root.iter(f"{{{ns_uri}}}comment"):
        author = c.get(f"{{{ns_uri}}}author") or ""
        texts = [t.text for t in c.iter(f"{{{ns_uri}}}t") if t.text]
        content = "".join(texts).strip()
        if content:
            memos.append(ExtractedMemo(
                location_context=f"댓글 ({author})" if author else "댓글",
                memo_content=content,
            ))
    return memos


# ── HWP (pyhwp / hwp5) ───────────────────────────────────────────────

def _extract_hwp(data: bytes, *, title: str, filename: str) -> UnifiedProofExtract:
    if not _has_module("hwp5"):
        raise ValueError(
            "한글(.hwp) 파일을 읽으려면 pyhwp가 필요합니다. "
            "터미널에서 `pip install --pre pyhwp six olefile` 후 다시 시도하거나, "
            "한글에서 HWPX/DOCX로 저장해 주세요."
        )
    try:
        return _extract_hwp_pyhwp(data, title=title)
    except ValueError:
        raise
    except Exception as error:  # noqa: BLE001
        raise ValueError(
            f"HWP 파일을 해석하지 못했습니다: {error}. "
            "손상된 파일이거나 암호/배포용 문서일 수 있습니다. HWPX·DOCX로 저장해 주세요."
        ) from error


def _extract_hwp_pyhwp(data: bytes, *, title: str) -> UnifiedProofExtract:
    """Use pyhwp (hwp5) TextTransform + model walk for memos."""
    from hwp5.xmlmodel import Hwp5File  # type: ignore
    from hwp5.hwp5txt import TextTransform  # type: ignore
    from hwp5.binmodel import ParaText  # type: ignore
    from hwp5.binmodel.controlchar import ControlChar  # type: ignore

    with tempfile.NamedTemporaryFile(suffix=".hwp", delete=False) as tmp:
        tmp.write(data)
        tmp_path = tmp.name
    out_path = tmp_path + ".txt"

    body = ""
    memos: list[ExtractedMemo] = []
    warnings: list[str] = []
    last_text = ""

    try:
        # 1) Preferred: official text transform (XSL)
        try:
            transform = TextTransform().transform_hwp5_to_text
            hwp5file = Hwp5File(tmp_path)
            try:
                with open(out_path, "wb") as dest:
                    transform(hwp5file, dest)
            finally:
                # Hwp5File is not a context manager on all versions
                closer = getattr(hwp5file, "close", None)
                if callable(closer):
                    try:
                        closer()
                    except Exception:
                        pass
            raw_out = Path(out_path).read_bytes()
            for enc in ("utf-8", "utf-16", "cp949"):
                try:
                    body = raw_out.decode(enc)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                body = raw_out.decode("utf-8", errors="replace")
            body = document_import.normalise_whitespace(body)
        except Exception as xsl_error:  # noqa: BLE001
            warnings.append(f"HWP 텍스트 변환기 실패, 레코드 파서로 재시도: {xsl_error}")
            body = ""

        # 2) Model walk — fill body if empty + harvest memo controls
        try:
            hwp = Hwp5File(tmp_path)
            try:
                paragraphs: list[str] = []
                bodytext = hwp.bodytext
                indexes = list(bodytext.section_indexes())
                for idx in indexes:
                    section = bodytext.section(idx)
                    for model in section.models():
                        if not isinstance(model, dict):
                            continue
                        mtype = model.get("type")
                        type_name = getattr(mtype, "__name__", str(mtype or ""))
                        content = model.get("content") or {}

                        if mtype is ParaText or type_name == "ParaText":
                            chunks = content.get("chunks") if isinstance(content, dict) else None
                            para_bits: list[str] = []
                            for item in chunks or []:
                                payload = item[1] if isinstance(item, (list, tuple)) and len(item) >= 2 else item
                                if isinstance(payload, str):
                                    para_bits.append(payload)
                                elif isinstance(payload, ControlChar):
                                    code = getattr(payload, "code", None)
                                    if code in (10, 13) or getattr(payload, "char", "") in ("\n", "\r"):
                                        para_bits.append("\n")
                            para = "".join(para_bits).replace("\r", "\n").strip()
                            if para:
                                paragraphs.append(para)
                                last_text = para
                        elif any(k in type_name for k in ("Memo", "Comment", "HiddenComment")):
                            raw = ""
                            if isinstance(content, dict):
                                raw = str(
                                    content.get("text")
                                    or content.get("memo")
                                    or content.get("chunks")
                                    or ""
                                )
                            raw = re.sub(r"\s+", " ", raw).strip()
                            if 1 < len(raw) < 500 and "object at" not in raw:
                                memos.append(ExtractedMemo(
                                    location_context=last_text[:80] or "(HWP 메모)",
                                    memo_content=raw,
                                ))

                if not body.strip() and paragraphs:
                    body = "\n\n".join(paragraphs)
                if not body.strip():
                    try:
                        prv = hwp.preview_text
                        preview = prv.text if hasattr(prv, "text") else str(prv)
                        if str(preview).strip():
                            body = str(preview).strip()
                            warnings.append("본문 대신 HWP 미리보기 텍스트를 사용했습니다.")
                    except Exception:
                        pass
            finally:
                closer = getattr(hwp, "close", None)
                if callable(closer):
                    try:
                        closer()
                    except Exception:
                        pass
        except Exception as model_error:  # noqa: BLE001
            if not body.strip():
                raise model_error
            warnings.append(f"메모 스캔 일부 실패: {model_error}")
    finally:
        for path in (tmp_path, out_path):
            try:
                Path(path).unlink(missing_ok=True)
            except Exception:
                pass

    if not body.strip():
        raise ValueError("HWP에서 본문 텍스트를 찾지 못했습니다.")

    body2, memos2 = _merge_text_memos(body, memos)
    return UnifiedProofExtract(
        text=body2,
        memos=memos2,
        format="hwp",
        parser="pyhwp",
        title=title,
        warnings=warnings,
    )
