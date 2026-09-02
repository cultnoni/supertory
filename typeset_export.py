"""Platform typeset presets (Munpia / KakaoPage / Ridibooks / Naver Series) and DOCX/HWPX export."""

from __future__ import annotations

import io
import re
import zipfile

from document_export import (
    ExportFile,
    _hwpx_section_xml_from_lines,
    _hwpx_skeleton_path,
    safe_download_name,
    validate_hwpx_package,
)

PLATFORM_ORDER = ("munpia", "kakaopage", "ridibooks", "naver_series")

_DEFAULT_FIELDS: dict = {
    "label": "",
    "font_family": "바탕체",
    "font_size_pt": 10,
    "line_height_percent": 150,
    "letter_spacing_pt": 0,
    "paragraph_indent_pt": 0,
    "paragraph_spacing_pt": 0,
    "margin_left_mm": 20,
    "margin_right_mm": 20,
    "margin_top_mm": 20,
    "margin_bottom_mm": 20,
    "mobile_viewport_px": 360,
    "is_verified": False,
    "is_default": False,
}

DEFAULT_PRESETS: dict[str, dict] = {
    "munpia": {
        **_DEFAULT_FIELDS,
        "label": "문피아",
        "line_height_percent": 140,
        "is_verified": True,
        "is_default": True,
    },
    "kakaopage": {
        **_DEFAULT_FIELDS,
        "label": "카카오페이지",
        "line_height_percent": 150,
        "is_verified": False,
        "is_default": True,
    },
    "ridibooks": {
        **_DEFAULT_FIELDS,
        "label": "리디북스",
        "line_height_percent": 160,
        "paragraph_indent_pt": 100,
        "is_verified": False,
        "is_default": True,
    },
    "naver_series": {
        **_DEFAULT_FIELDS,
        "label": "네이버 시리즈",
        "line_height_percent": 150,
        "is_verified": False,
        "is_default": True,
    },
}

_INT_FIELDS = frozenset({
    "font_size_pt",
    "line_height_percent",
    "paragraph_indent_pt",
    "paragraph_spacing_pt",
    "margin_left_mm",
    "margin_right_mm",
    "margin_top_mm",
    "margin_bottom_mm",
    "mobile_viewport_px",
})
_FLOAT_FIELDS = frozenset({
    "letter_spacing_pt",
})

_FIELD_RANGE = {
    "font_size_pt": (1, 72),
    "line_height_percent": (50, 400),
    "letter_spacing_pt": (-20, 40),
    "paragraph_indent_pt": (0, 400),
    "paragraph_spacing_pt": (0, 120),
    "margin_left_mm": (0, 80),
    "margin_right_mm": (0, 80),
    "margin_top_mm": (0, 80),
    "margin_bottom_mm": (0, 80),
    "mobile_viewport_px": (200, 800),
}

_FONT_FILE_NAMES = {
    "바탕체": "Batang",
    "바탕": "Batang",
    "돋움": "Dotum",
    "돋움체": "Dotum",
    "굴림": "Gulim",
    "굴림체": "Gulim",
    "맑은 고딕": "Malgun Gothic",
    "맑은고딕": "Malgun Gothic",
}

_HANGUL_FONT_NAMES = {
    "바탕체": "바탕",
    "바탕": "바탕",
    "Batang": "바탕",
    "돋움체": "돋움",
    "돋움": "돋움",
    "Dotum": "돋움",
    "굴림체": "굴림",
    "굴림": "굴림",
    "Gulim": "굴림",
    "맑은 고딕": "맑은 고딕",
    "맑은고딕": "맑은 고딕",
    "Malgun Gothic": "맑은 고딕",
}

_HWPUNIT_PER_MM = 7200 / 25.4
_TYPESET_EXPORT_FORMATS = frozenset({"docx", "hwpx"})


def is_builtin_platform(platform_id: str) -> bool:
    return str(platform_id or "").strip() in PLATFORM_ORDER


def map_font_family(name: str) -> str:
    text = str(name or "").strip() or "바탕체"
    return _FONT_FILE_NAMES.get(text, text)


def map_hangul_font_family(name: str) -> str:
    text = str(name or "").strip() or "바탕체"
    return _HANGUL_FONT_NAMES.get(text, text)


def format_key_from_body(body: dict | None) -> str:
    raw = body if isinstance(body, dict) else {}
    key = str(raw.get("format") or raw.get("format_key") or "docx").strip().lower()
    if key in {"doc", "word"}:
        key = "docx"
    if key in {"hwp", "hangul", "hwpml"}:
        key = "hwpx"
    if key not in _TYPESET_EXPORT_FORMATS:
        raise ValueError("조판 내보내기는 Word(.docx)와 한글(.hwpx)만 지원해요.")
    return key


def split_typeset_paragraphs(plain_text: str) -> list[str]:
    """Split manuscript body on line breaks (platform paragraph unit)."""
    text = str(plain_text or "").replace("\r\n", "\n").replace("\r", "\n")
    text = text.strip("\n")
    if not text.strip():
        return [""]
    return text.split("\n")


def _as_int(value, default: int) -> int:
    try:
        if isinstance(value, bool):
            raise TypeError
        return int(round(float(value)))
    except (TypeError, ValueError):
        return int(default)


def _as_float(value, default: float) -> float:
    try:
        if isinstance(value, bool):
            raise TypeError
        return float(value)
    except (TypeError, ValueError):
        return float(default)


def _as_bool(value, default: bool) -> bool:
    if isinstance(value, bool):
        return value
    if isinstance(value, (int, float)) and not isinstance(value, bool):
        return bool(value)
    text = str(value or "").strip().lower()
    if text in {"1", "true", "yes", "y", "on"}:
        return True
    if text in {"0", "false", "no", "n", "off", ""}:
        return False
    return bool(default)


def _clamp(name: str, value: int) -> int:
    bounds = _FIELD_RANGE.get(name)
    if not bounds:
        return value
    low, high = bounds
    return max(int(low), min(int(high), int(value)))


def _clamp_float(name: str, value: float) -> float:
    bounds = _FIELD_RANGE.get(name)
    if not bounds:
        return value
    low, high = bounds
    return max(float(low), min(float(high), float(value)))


def normalize_preset(
    raw: dict | None,
    *,
    fallback: dict | None = None,
    platform_id: str = "",
) -> dict:
    base = {**_DEFAULT_FIELDS, **(fallback or {})}
    src = raw if isinstance(raw, dict) else {}
    out = dict(base)
    if "label" in src:
        out["label"] = str(src.get("label") or base.get("label") or "").strip()
    if "font_family" in src:
        family = str(src.get("font_family") or "").strip()
        out["font_family"] = family or str(base.get("font_family") or "바탕체")
    for key in _INT_FIELDS:
        if key in src:
            out[key] = _clamp(key, _as_int(src.get(key), int(base.get(key) or 0)))
    for key in _FLOAT_FIELDS:
        if key in src:
            out[key] = _clamp_float(key, _as_float(src.get(key), float(base.get(key) or 0)))
    if "is_verified" in src:
        out["is_verified"] = _as_bool(src.get("is_verified"), bool(base.get("is_verified")))
    out["label"] = str(out.get("label") or "").strip()
    out["font_family"] = str(out.get("font_family") or "바탕체").strip() or "바탕체"
    for key in _INT_FIELDS:
        out[key] = _clamp(key, _as_int(out.get(key), int(_DEFAULT_FIELDS[key])))
    for key in _FLOAT_FIELDS:
        out[key] = _clamp_float(key, _as_float(out.get(key), float(_DEFAULT_FIELDS[key])))
    out["is_verified"] = bool(out.get("is_verified"))
    out["is_default"] = is_builtin_platform(platform_id)
    return out


def _require_docx():
    try:
        from docx import Document  # type: ignore
        from docx.oxml import OxmlElement  # type: ignore
        from docx.oxml.ns import qn  # type: ignore
        from docx.shared import Mm, Pt  # type: ignore
    except ImportError as error:
        raise ValueError(
            "python-docx가 설치되어 있지 않아요. "
            "터미널에서 pip install python-docx 후 SuperTORY를 다시 실행해 주세요."
        ) from error
    return Document, OxmlElement, qn, Mm, Pt


def _set_run_rfonts(run, mapped: str, east_asia: str, qn, OxmlElement) -> None:
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), mapped)
    rFonts.set(qn("w:hAnsi"), mapped)
    rFonts.set(qn("w:cs"), mapped)
    rFonts.set(qn("w:eastAsia"), east_asia)


def _letter_spacing_is_relative(value: float) -> bool:
    """Dropdown 자간(-0.05~0.1) is em; legacy integers stay pt."""
    return 0 < abs(float(value or 0)) <= 0.5


def _set_run_letter_spacing(run, letter_spacing_pt: float, font_size_pt: float, qn, OxmlElement) -> None:
    raw = float(letter_spacing_pt or 0)
    if not raw:
        return
    size = max(1.0, float(font_size_pt or 10))
    twips = raw * size * 20 if _letter_spacing_is_relative(raw) else raw * 20
    rPr = run._element.get_or_add_rPr()
    spacing = rPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        rPr.append(spacing)
    spacing.set(qn("w:val"), str(int(round(twips))))


def _apply_paragraph_format(paragraph, preset: dict, Pt) -> None:
    pf = paragraph.paragraph_format
    pf.line_spacing = max(0.5, float(preset["line_height_percent"]) / 100.0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(int(preset["paragraph_spacing_pt"]))
    pf.first_line_indent = Pt(int(preset["paragraph_indent_pt"]))
    pf.widow_control = True


def _apply_run_style(run, preset: dict, *, mapped: str, east_asia: str, Pt, qn, OxmlElement) -> None:
    run.font.name = mapped
    run.font.size = Pt(int(preset["font_size_pt"]))
    _set_run_rfonts(run, mapped, east_asia, qn, OxmlElement)
    _set_run_letter_spacing(
        run,
        float(preset["letter_spacing_pt"] or 0),
        float(preset["font_size_pt"] or 10),
        qn,
        OxmlElement,
    )


def build_typeset_docx(plain_text: str, preset: dict) -> bytes:
    """Build a DOCX whose font, line height, indent, and margins follow the preset."""
    Document, OxmlElement, qn, Mm, Pt = _require_docx()
    preset = normalize_preset(preset)
    family = str(preset["font_family"])
    mapped = map_font_family(family)
    east_asia = family if family in {"바탕", "바탕체"} else mapped
    if east_asia == "바탕체":
        east_asia = "바탕"

    document = Document()
    section = document.sections[0]
    section.left_margin = Mm(int(preset["margin_left_mm"]))
    section.right_margin = Mm(int(preset["margin_right_mm"]))
    section.top_margin = Mm(int(preset["margin_top_mm"]))
    section.bottom_margin = Mm(int(preset["margin_bottom_mm"]))

    style = document.styles["Normal"]
    style.font.name = mapped
    style.font.size = Pt(int(preset["font_size_pt"]))
    try:
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:ascii"), mapped)
        rFonts.set(qn("w:hAnsi"), mapped)
        rFonts.set(qn("w:cs"), mapped)
        rFonts.set(qn("w:eastAsia"), east_asia)
    except Exception:
        pass
    style.paragraph_format.line_spacing = max(0.5, float(preset["line_height_percent"]) / 100.0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(int(preset["paragraph_spacing_pt"]))
    style.paragraph_format.first_line_indent = Pt(int(preset["paragraph_indent_pt"]))

    lines = split_typeset_paragraphs(plain_text)
    for line in lines:
        paragraph = document.add_paragraph()
        _apply_paragraph_format(paragraph, preset, Pt)
        run = paragraph.add_run(line)
        _apply_run_style(
            run,
            preset,
            mapped=mapped,
            east_asia=east_asia,
            Pt=Pt,
            qn=qn,
            OxmlElement=OxmlElement,
        )

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _hwp_units_from_mm(mm: int) -> int:
    return int(round(float(mm) * _HWPUNIT_PER_MM))


def _hwp_units_from_pt(pt: int) -> int:
    return int(pt) * 100


def _xml_attr_escape(text: str) -> str:
    return (
        str(text or "")
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
        .replace("'", "&apos;")
    )


def _patch_hwpx_header(header: str, preset: dict) -> str:
    font_name = _xml_attr_escape(map_hangul_font_family(str(preset.get("font_family") or "바탕체")))
    height = max(100, int(preset["font_size_pt"]) * 100)
    size_pt = max(1, int(preset["font_size_pt"]))
    raw_spacing = float(preset["letter_spacing_pt"] or 0)
    if _letter_spacing_is_relative(raw_spacing):
        spacing = int(round(raw_spacing * 100.0))
    else:
        spacing = int(round(raw_spacing / float(size_pt) * 100.0))
    spacing = max(-50, min(50, spacing))
    indent_hwp = max(0, _hwp_units_from_pt(int(preset["paragraph_indent_pt"])))
    indent_char = indent_hwp // 2
    after_hwp = max(0, _hwp_units_from_pt(int(preset["paragraph_spacing_pt"])))
    after_char = after_hwp // 2
    line_percent = max(50, min(400, int(preset["line_height_percent"])))

    patched = header.replace('face="함초롬바탕"', f'face="{font_name}"')

    def _replace_tagged(xml: str, open_tag: str, close_tag: str, replacer) -> str:
        start = xml.find(open_tag)
        if start < 0:
            return xml
        end = xml.find(close_tag, start)
        if end < 0:
            return xml
        end += len(close_tag)
        return xml[:start] + replacer(xml[start:end]) + xml[end:]

    def patch_char(block: str) -> str:
        block = re.sub(r'\bheight="\d+"', f'height="{height}"', block, count=1)
        block = re.sub(
            r"<hh:spacing hangul=\"[^\"]*\" latin=\"[^\"]*\" hanja=\"[^\"]*\" "
            r"japanese=\"[^\"]*\" other=\"[^\"]*\" symbol=\"[^\"]*\" user=\"[^\"]*\"/>",
            (
                f'<hh:spacing hangul="{spacing}" latin="{spacing}" hanja="{spacing}" '
                f'japanese="{spacing}" other="{spacing}" symbol="{spacing}" user="{spacing}"/>'
            ),
            block,
            count=1,
        )
        return block

    def patch_margin_block(inner: str, *, indent: int, after: int) -> str:
        inner = re.sub(
            r'(<hc:intent value=")[^"]+(")',
            lambda match: f"{match.group(1)}{indent}{match.group(2)}",
            inner,
            count=1,
        )
        inner = re.sub(
            r'(<hc:next value=")[^"]+(")',
            lambda match: f"{match.group(1)}{after}{match.group(2)}",
            inner,
            count=1,
        )
        return inner

    def patch_para(block: str) -> str:
        block = re.sub(
            r'(<hh:lineSpacing type="PERCENT" value=")[^"]+(")',
            lambda match: f"{match.group(1)}{line_percent}{match.group(2)}",
            block,
        )
        block = _replace_tagged(
            block,
            "<hp:case ",
            "</hp:case>",
            lambda inner: patch_margin_block(inner, indent=indent_char, after=after_char),
        )
        return _replace_tagged(
            block,
            "<hp:default>",
            "</hp:default>",
            lambda inner: patch_margin_block(inner, indent=indent_hwp, after=after_hwp),
        )

    patched = _replace_tagged(patched, '<hh:charPr id="0"', "</hh:charPr>", patch_char)
    return _replace_tagged(patched, '<hh:paraPr id="0"', "</hh:paraPr>", patch_para)


def _patch_hwpx_page_margins(section: str, preset: dict) -> str:
    left = _hwp_units_from_mm(int(preset["margin_left_mm"]))
    right = _hwp_units_from_mm(int(preset["margin_right_mm"]))
    top = _hwp_units_from_mm(int(preset["margin_top_mm"]))
    bottom = _hwp_units_from_mm(int(preset["margin_bottom_mm"]))

    def repl(match: re.Match[str]) -> str:
        return (
            f'{match.group(1)}{match.group(2)}{match.group(3)}{match.group(4)}'
            f'{match.group(5)}{match.group(6)}{match.group(7)}{left}'
            f'{match.group(8)}{right}{match.group(9)}{top}{match.group(10)}{bottom}{match.group(11)}'
        )

    return re.sub(
        r'(<hp:margin header=")([^"]+)(" footer=")([^"]+)(" gutter=")([^"]+)'
        r'(" left=")[^"]+(" right=")[^"]+(" top=")[^"]+(" bottom=")[^"]+("/>)',
        repl,
        section,
        count=1,
    )


def build_typeset_hwpx(plain_text: str, preset: dict) -> bytes:
    """Build an HWPX whose font, line height, indent, and margins follow the preset."""
    preset = normalize_preset(preset)
    skeleton = _hwpx_skeleton_path()
    if not skeleton.is_file():
        raise ValueError(
            "한글(HWPX) 내보내기 템플릿(assets/hwpx_skeleton.hwpx)을 찾지 못했습니다."
        )

    with zipfile.ZipFile(skeleton, "r") as src:
        header = _patch_hwpx_header(
            src.read("Contents/header.xml").decode("utf-8"),
            preset,
        )
        skeleton_section = _patch_hwpx_page_margins(
            src.read("Contents/section0.xml").decode("utf-8"),
            preset,
        )
        section_xml = _hwpx_section_xml_from_lines(
            split_typeset_paragraphs(plain_text),
            skeleton_section,
        )
        preview = str(plain_text or "")[:4000].encode("utf-8")
        buffer = io.BytesIO()
        ordered = ["mimetype"] + [name for name in src.namelist() if name != "mimetype"]
        with zipfile.ZipFile(buffer, "w") as dst:
            for name in ordered:
                raw = src.read(name)
                if name == "Contents/header.xml":
                    raw = header.encode("utf-8")
                elif name == "Contents/section0.xml":
                    raw = section_xml
                elif name == "Preview/PrvText.txt":
                    raw = preview
                compress = (
                    zipfile.ZIP_STORED
                    if name == "mimetype"
                    else src.getinfo(name).compress_type
                )
                info = zipfile.ZipInfo(filename=name)
                info.compress_type = compress
                if name == "mimetype":
                    info.flag_bits = 0
                dst.writestr(info, raw)
        data = buffer.getvalue()

    problems = validate_hwpx_package(data)
    if problems:
        raise ValueError("한글(HWPX) 패키지 생성에 실패했습니다: " + "; ".join(problems))
    return data


def export_typeset_file(
    *,
    plain_text: str,
    preset: dict,
    scene_title: str,
    platform_id: str,
    format_key: str = "docx",
) -> ExportFile:
    key = format_key_from_body({"format": format_key})
    if key == "hwpx":
        data = build_typeset_hwpx(plain_text, preset)
        ext = ".hwpx"
        mime = "application/hwp+zip"
    else:
        data = build_typeset_docx(plain_text, preset)
        ext = ".docx"
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    label = str(preset.get("label") or platform_id or "조판").strip() or "조판"
    title = str(scene_title or "").strip() or "회차"
    filename = safe_download_name(f"{title}_{label}", ext)
    return ExportFile(
        filename=filename,
        mime=mime,
        data=data,
        format_key=key,
    )


def platform_id_from_body(body: dict | None) -> str:
    raw = body if isinstance(body, dict) else {}
    return str(raw.get("platform_id") or raw.get("platform") or "").strip()


def scene_id_from_body(body: dict | None) -> int:
    raw = body if isinstance(body, dict) else {}
    value = raw.get("scene_id")
    if value is None:
        value = raw.get("chapter_id")
    if value is None:
        value = raw.get("episode_id")
    try:
        scene_id = int(value)
    except (TypeError, ValueError) as error:
        raise ValueError("회차를 선택해 주세요.") from error
    if scene_id <= 0:
        raise ValueError("회차를 선택해 주세요.")
    return scene_id
