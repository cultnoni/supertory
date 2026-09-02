"""Platform typeset presets API and DOCX export."""

from __future__ import annotations

import io
import json
import tempfile
import threading
import unittest
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

import app
import document_export
import typeset_export
from repositories.typeset_repository import TypesetRepository
from services.typeset_service import TypesetService


class TypesetPresetUnitTests(unittest.TestCase):
    def test_split_paragraphs_on_newlines(self) -> None:
        self.assertEqual(
            typeset_export.split_typeset_paragraphs("첫 줄\n둘째 줄\n\n넷째"),
            ["첫 줄", "둘째 줄", "", "넷째"],
        )

    def test_docx_applies_font_line_height_indent_and_margins(self) -> None:
        preset = typeset_export.normalize_preset({
            "label": "리디북스",
            "font_family": "바탕체",
            "font_size_pt": 10,
            "line_height_percent": 160,
            "letter_spacing_pt": 0,
            "paragraph_indent_pt": 100,
            "paragraph_spacing_pt": 0,
            "margin_left_mm": 20,
            "margin_right_mm": 20,
            "margin_top_mm": 20,
            "margin_bottom_mm": 20,
            "mobile_viewport_px": 360,
        })
        data = typeset_export.build_typeset_docx("바람이 불었다.\n문이 열렸다.", preset)
        self.assertGreater(len(data), 100)
        self.assertEqual(data[:2], b"PK")

        document = Document(io.BytesIO(data))
        section = document.sections[0]
        self.assertAlmostEqual(section.left_margin.mm, 20, delta=0.2)
        self.assertAlmostEqual(section.right_margin.mm, 20, delta=0.2)
        self.assertAlmostEqual(section.top_margin.mm, 20, delta=0.2)
        self.assertAlmostEqual(section.bottom_margin.mm, 20, delta=0.2)

        texts = [para.text for para in document.paragraphs]
        self.assertEqual(texts[:2], ["바람이 불었다.", "문이 열렸다."])
        first = document.paragraphs[0]
        self.assertAlmostEqual(float(first.paragraph_format.line_spacing), 1.6, places=2)
        self.assertEqual(first.paragraph_format.first_line_indent, Pt(100))
        run = first.runs[0]
        self.assertEqual(run.font.size, Pt(10))
        rFonts = run._element.find(qn("w:rPr")).find(qn("w:rFonts"))
        self.assertIsNotNone(rFonts)
        self.assertIn(rFonts.get(qn("w:eastAsia")), {"바탕", "바탕체", "Batang"})

        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            xml = archive.read("word/document.xml").decode("utf-8")
        self.assertIn("바람이 불었다.", xml)
        self.assertIn("문이 열렸다.", xml)

    def test_hwpx_applies_font_line_height_indent_and_margins(self) -> None:
        preset = typeset_export.normalize_preset({
            "label": "리디북스",
            "font_family": "바탕체",
            "font_size_pt": 10,
            "line_height_percent": 160,
            "letter_spacing_pt": 0,
            "paragraph_indent_pt": 100,
            "paragraph_spacing_pt": 0,
            "margin_left_mm": 20,
            "margin_right_mm": 20,
            "margin_top_mm": 20,
            "margin_bottom_mm": 20,
            "mobile_viewport_px": 360,
        })
        data = typeset_export.build_typeset_hwpx("바람이 불었다.\n문이 열렸다.", preset)
        self.assertEqual(document_export.validate_hwpx_package(data), [])
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            header = archive.read("Contents/header.xml").decode("utf-8")
            section = archive.read("Contents/section0.xml").decode("utf-8")
        self.assertIn('face="바탕"', header)
        self.assertIn('type="PERCENT" value="160"', header)
        self.assertIn('intent value="10000"', header)
        self.assertIn("바람이 불었다.", section)
        self.assertIn("문이 열렸다.", section)
        self.assertIn(f'left="{typeset_export._hwp_units_from_mm(20)}"', section)

    def test_update_rejects_unknown_platform(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            service = TypesetService(TypesetRepository(
                root=Path(app.ROOT),
                data_dir=Path(folder),
            ))
            with self.assertRaises(ValueError):
                service.update_preset("unknown_site", {"font_size_pt": 12})

    def test_letter_spacing_preserves_fractional_values(self) -> None:
        preset = typeset_export.normalize_preset({
            "letter_spacing_pt": 0.05,
            "font_size_pt": 10,
        })
        self.assertAlmostEqual(preset["letter_spacing_pt"], 0.05)
        ridibooks = typeset_export.normalize_preset(
            typeset_export.DEFAULT_PRESETS["ridibooks"],
            platform_id="ridibooks",
        )
        self.assertEqual(ridibooks["paragraph_indent_pt"], 100)

    def test_service_creates_copies_updates_and_deletes_custom_preset(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            data_dir = Path(folder)
            repository = TypesetRepository(root=Path(app.ROOT), data_dir=data_dir)
            service = TypesetService(repository)

            created = service.create_preset("내 조판", copy_from="ridibooks")
            platform_id = str(created["platform_id"])
            self.assertEqual(platform_id, "nae_jopan")
            self.assertEqual(created["preset"]["paragraph_indent_pt"], 100)
            self.assertFalse(created["preset"]["is_default"])
            self.assertFalse(created["preset"]["is_verified"])

            updated = service.update_preset(platform_id, {
                "label": "내 조판 수정",
                "font_size_pt": 13,
            })
            self.assertEqual(updated["preset"]["label"], "내 조판 수정")
            self.assertEqual(updated["preset"]["font_size_pt"], 13)
            self.assertTrue((data_dir / "typeset_presets.json").is_file())

            deleted = service.delete_preset(platform_id)
            self.assertNotIn(platform_id, deleted["presets"])

    def test_service_rejects_deleting_default_preset(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            service = TypesetService(TypesetRepository(
                root=Path(app.ROOT),
                data_dir=Path(folder),
            ))
            with self.assertRaisesRegex(
                ValueError,
                "기본 조판양식은 삭제할 수 없습니다",
            ):
                service.delete_preset("munpia")

    def test_repository_crud_contract_and_corrupt_runtime_fallback(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            data_dir = Path(folder)
            repository = TypesetRepository(root=Path(app.ROOT), data_dir=data_dir)
            self.assertEqual(len(repository.list_presets(None)), 4)
            self.assertEqual(repository.get_preset(None, "munpia")["label"], "문피아")

            created = repository.create_preset(None, "저장소 검증", {
                "platform_id": "repository_check",
                "font_size_pt": 12,
                "is_default": False,
            })
            self.assertEqual(created["font_size_pt"], 12)
            updated = repository.update_preset(
                None,
                "repository_check",
                {**created, "font_size_pt": 13},
            )
            self.assertEqual(updated["font_size_pt"], 13)
            self.assertTrue(repository.delete_preset(None, "repository_check"))
            self.assertFalse(repository.delete_preset(None, "repository_check"))

            (data_dir / "typeset_presets.json").write_text("{broken", encoding="utf-8")
            fallback = repository.list_presets(None)
            self.assertEqual(
                {row["platform_id"] for row in fallback},
                {"munpia", "kakaopage", "ridibooks", "naver_series"},
            )


class TypesetPresetApiTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temporary_directory = tempfile.TemporaryDirectory()
        self.original_data_dir = app.DATA_DIR
        self.original_database_path = app.DATABASE_PATH
        app.DATA_DIR = Path(self.temporary_directory.name) / "data"
        app.DATABASE_PATH = app.DATA_DIR / "supertory.sqlite3"
        app.initialise_database()
        self.server = app.ThreadingHTTPServer(("127.0.0.1", 0), app.SuperToryHandler)
        self.thread = threading.Thread(target=self.server.serve_forever, daemon=True)
        self.thread.start()

    def tearDown(self) -> None:
        self.server.shutdown()
        self.thread.join()
        self.server.server_close()
        app.DATA_DIR = self.original_data_dir
        app.DATABASE_PATH = self.original_database_path
        self.temporary_directory.cleanup()

    def request(
        self,
        method: str,
        path: str,
        payload: dict | None = None,
        *,
        json_response: bool = True,
    ) -> tuple[int, object]:
        import http.client

        connection = http.client.HTTPConnection("127.0.0.1", self.server.server_port)
        body = json.dumps(payload, ensure_ascii=False).encode("utf-8") if payload is not None else None
        headers = {"Content-Type": "application/json"} if body else {}
        connection.request(method, path, body, headers)
        response = connection.getresponse()
        raw = response.read()
        connection.close()
        if json_response:
            return response.status, json.loads(raw.decode("utf-8"))
        return response.status, raw

    def _make_scene(self) -> int:
        status, project = self.request(
            "POST", "/api/projects", {"title": "조판소설", "main_genre": "판타지"}
        )
        self.assertEqual(status, 201)
        status, chapter = self.request(
            "POST", f"/api/projects/{project['id']}/chapters", {"title": "1장"}
        )
        self.assertEqual(status, 201)
        status, scene = self.request(
            "POST", f"/api/chapters/{chapter['id']}/scenes", {"title": "1화"}
        )
        self.assertEqual(status, 201)
        status, detail = self.request("GET", f"/api/scenes/{scene['id']}")
        self.assertEqual(status, 200)
        status, _ = self.request("PUT", f"/api/scenes/{scene['id']}", {
            "title": "1화",
            "status": "draft",
            "content_md": "첫 문단입니다.\n둘째 문단입니다.",
            "row_version": detail["row_version"],
        })
        self.assertEqual(status, 200)
        return int(scene["id"])

    def test_get_presets_includes_four_platforms(self) -> None:
        status, data = self.request("GET", "/api/typeset/presets")
        self.assertEqual(status, 200)
        presets = data["presets"]
        self.assertEqual(
            set(presets),
            {"munpia", "kakaopage", "ridibooks", "naver_series"},
        )
        self.assertTrue(presets["munpia"]["is_verified"])
        self.assertFalse(presets["kakaopage"]["is_verified"])
        self.assertEqual(presets["munpia"]["line_height_percent"], 140)
        self.assertEqual(presets["ridibooks"]["paragraph_indent_pt"], 100)

    def test_put_updates_runtime_file_not_seed(self) -> None:
        status, data = self.request(
            "PUT",
            "/api/typeset/presets/kakaopage",
            {"font_size_pt": 11, "line_height_percent": 155, "mobile_viewport_px": 390},
        )
        self.assertEqual(status, 200)
        self.assertEqual(data["preset"]["font_size_pt"], 11)
        self.assertEqual(data["preset"]["line_height_percent"], 155)
        self.assertEqual(data["preset"]["mobile_viewport_px"], 390)
        self.assertFalse(data["preset"]["is_verified"])

        runtime = Path(app.DATA_DIR) / "typeset_presets.json"
        self.assertTrue(runtime.is_file())
        saved = json.loads(runtime.read_text(encoding="utf-8"))
        self.assertEqual(saved["kakaopage"]["font_size_pt"], 11)

        seed = Path(app.ROOT) / "data" / "typeset_presets.json"
        seed_data = json.loads(seed.read_text(encoding="utf-8"))
        self.assertEqual(seed_data["kakaopage"]["font_size_pt"], 10)

        status, again = self.request("GET", "/api/typeset/presets")
        self.assertEqual(status, 200)
        self.assertEqual(again["presets"]["kakaopage"]["font_size_pt"], 11)

    def test_custom_preset_create_rename_delete_round_trip(self) -> None:
        status, created = self.request("POST", "/api/typeset/presets", {
            "label": "투고용",
            "copy_from": "ridibooks",
        })
        self.assertEqual(status, 200, created)
        platform_id = str(created["platform_id"])
        self.assertFalse(created["preset"]["is_default"])
        self.assertEqual(created["preset"]["paragraph_indent_pt"], 100)

        status, updated = self.request(
            "PUT",
            f"/api/typeset/presets/{platform_id}",
            {"label": "투고용 수정", "font_size_pt": 12},
        )
        self.assertEqual(status, 200, updated)
        self.assertEqual(updated["preset"]["label"], "투고용 수정")
        self.assertEqual(updated["preset"]["font_size_pt"], 12)

        status, deleted = self.request(
            "DELETE",
            f"/api/typeset/presets/{platform_id}",
        )
        self.assertEqual(status, 200, deleted)
        self.assertNotIn(platform_id, deleted["presets"])

    def test_delete_rejects_default_preset(self) -> None:
        status, result = self.request("DELETE", "/api/typeset/presets/munpia")
        self.assertEqual(status, 404)
        self.assertEqual(result.get("error"), "기본 조판양식은 삭제할 수 없습니다")

    def test_export_docx_to_folder(self) -> None:
        scene_id = self._make_scene()
        export_dir = Path(app.DATA_DIR) / "typeset-out"
        status, result = self.request("POST", "/api/typeset/export", {
            "chapter_id": scene_id,
            "platform_id": "munpia",
            "save_to_folder": True,
            "export_dir": str(export_dir),
            "reveal_after_save": False,
        })
        self.assertEqual(status, 200, result)
        self.assertTrue(result.get("ok"))
        self.assertTrue(result.get("saved"))
        path = Path(result["path"])
        self.assertTrue(path.is_file())
        self.assertTrue(path.name.endswith(".docx"))
        document = Document(str(path))
        texts = [para.text for para in document.paragraphs]
        self.assertIn("첫 문단입니다.", texts)
        self.assertIn("둘째 문단입니다.", texts)
        first = document.paragraphs[0]
        self.assertAlmostEqual(float(first.paragraph_format.line_spacing), 1.4, places=2)
        self.assertEqual(first.runs[0].font.size, Pt(10))
        self.assertAlmostEqual(document.sections[0].left_margin, Mm(20), delta=Mm(0.2))

    def test_export_hwpx_to_folder(self) -> None:
        scene_id = self._make_scene()
        export_dir = Path(app.DATA_DIR) / "typeset-out"
        status, result = self.request("POST", "/api/typeset/export", {
            "chapter_id": scene_id,
            "platform_id": "munpia",
            "format": "hwpx",
            "save_to_folder": True,
            "export_dir": str(export_dir),
            "reveal_after_save": False,
        })
        self.assertEqual(status, 200, result)
        self.assertTrue(result.get("ok"))
        self.assertEqual(result.get("format"), "hwpx")
        path = Path(result["path"])
        self.assertTrue(path.is_file())
        self.assertTrue(path.name.endswith(".hwpx"))
        data = path.read_bytes()
        self.assertEqual(document_export.validate_hwpx_package(data), [])
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            header = archive.read("Contents/header.xml").decode("utf-8")
            section = archive.read("Contents/section0.xml").decode("utf-8")
        self.assertIn('face="바탕"', header)
        self.assertIn('height="1000"', header)
        self.assertIn('type="PERCENT" value="140"', header)
        self.assertIn("첫 문단입니다.", section)
        self.assertIn("둘째 문단입니다.", section)
        expected_margin = typeset_export._hwp_units_from_mm(20)
        self.assertIn(f'left="{expected_margin}"', section)
        self.assertIn(f'right="{expected_margin}"', section)

    def test_export_rejects_unknown_format(self) -> None:
        scene_id = self._make_scene()
        status, result = self.request("POST", "/api/typeset/export", {
            "chapter_id": scene_id,
            "platform_id": "munpia",
            "format": "pdf",
            "save_to_folder": False,
        })
        self.assertEqual(status, 400)
        self.assertIn("docx", str(result.get("error") or "").lower())

    def test_export_requires_scene(self) -> None:
        status, result = self.request("POST", "/api/typeset/export", {
            "platform_id": "munpia",
            "save_to_folder": False,
        })
        self.assertEqual(status, 400)
        self.assertIn("회차", str(result.get("error") or ""))
